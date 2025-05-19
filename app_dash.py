import dash
from dash import html, dcc, Input, Output, State, callback
import dash_bootstrap_components as dbc
import plotly.graph_objects as go
import pandas as pd
import fitz  # pymupdf
import re
from docx import Document
from num2words import num2words
from datetime import datetime, timedelta
import calendar
import io
import base64
import os
import json
from resources import CODIGOS_BRUTO, CODIGOS_DEDUCCIONES, MOTIVOS, TOPE_MAXIMO_PRESTAMO, TASA_ANUAL
import logging
import uuid
import tempfile
from flask import send_file
from docx.shared import Pt

# Configuración de logging
def setup_logging():
    # Obtener el entorno (desarrollo o producción)
    is_production = os.environ.get('RENDER', 'false').lower() == 'true'
    
    # Configurar el logger principal
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    
    # Formato común para todos los logs
    formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
    
    # En desarrollo, escribir a archivos
    if not is_production:
        # Logs técnicos
        log_filename = 'logs.txt'
        file_handler = logging.FileHandler(log_filename)
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)
        
        # Logs de usuario
        user_logger = logging.getLogger('user')
        user_logger.setLevel(logging.INFO)
        user_handler = logging.FileHandler('user_log.txt')
        user_handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s'))
        user_logger.addHandler(user_handler)
        
        # Logs de métricas
        metrics_logger = logging.getLogger('metrics')
        metrics_logger.setLevel(logging.INFO)
        metrics_handler = logging.FileHandler('metrics.json')
        metrics_handler.setFormatter(logging.Formatter('%(message)s'))
        metrics_logger.addHandler(metrics_handler)
    
    # En ambos entornos, escribir a consola
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    root_logger.addHandler(console_handler)

# Inicializar logging
setup_logging()

def log_user_action(action, details):
    """Función auxiliar para registrar acciones del usuario"""
    user_logger = logging.getLogger('user')
    user_logger.info(f"{action}: {details}")

def log_metric(event_type, data):
    """Función auxiliar para registrar métricas en formato JSON"""
    metric_data = {
        'timestamp': datetime.now().isoformat(),
        'event_type': event_type,
        'data': data
    }
    metrics_logger = logging.getLogger('metrics')
    metrics_logger.info(json.dumps(metric_data))

# Inicializar la aplicación Dash
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])

# Estilos CSS personalizados
app.index_string = '''
<!DOCTYPE html>
<html>
    <head>
        {%metas%}
        <title>Simulador de Adelanto de Haberes</title>
        {%favicon%}
        {%css%}
        <style>
            :root {
                --primary-color: #1f4e79;
                --secondary-color: #2c6aa0;
                --background-color: #f0f2f6;
                --text-color: #333333;
                --white: #ffffff;
            }

            body {
                background-color: var(--background-color);
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            }

            .navbar {
                background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
                padding: 2rem 2rem;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }

            .navbar-title {
                text-align: center;
                width: 100%;
            }

            .navbar-title h3 {
                color: white !important;
                font-size: 3.2rem;
                font-weight: 900;
                letter-spacing: 0.5px;
                margin: 0;
                text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
                text-transform: uppercase;
                -webkit-text-fill-color: white;
                -webkit-text-stroke: 1px white;
            }

            .navbar-subtitle {
                font-size: 0.9rem;
                color: rgba(255,255,255,0.9);
                margin-top: 0.2rem;
                font-weight: 400;
            }

            .navbar-logo {
                font-size: 2rem;
                color: var(--white);
                margin-right: 0.5rem;
            }

            .card {
                border: none;
                border-radius: 8px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                margin-bottom: 1.5rem;
            }

            .card-header {
                background-color: var(--primary-color);
                color: var(--white);
                font-weight: 600;
                border-radius: 8px 8px 0 0 !important;
            }

            .btn-primary {
                background-color: var(--primary-color);
                border: none;
                padding: 0.8rem 1.5rem;
                font-weight: 600;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                transition: all 0.3s ease;
            }

            .btn-primary:hover {
                background-color: var(--secondary-color);
                transform: translateY(-2px);
                box-shadow: 0 4px 8px rgba(0,0,0,0.2);
            }

            .form-control {
                border-radius: 6px;
                padding: 0.8rem;
                border: 1px solid #ced4da;
                box-shadow: inset 0 1px 3px rgba(0,0,0,0.05);
            }

            .form-control:focus {
                border-color: var(--primary-color);
                box-shadow: 0 0 0 0.2rem rgba(31, 78, 121, 0.25);
            }

            .alert {
                border-radius: 8px;
                border: none;
                box-shadow: 0 2px 4px rgba(0,0,0,0.05);
            }

            .alert-success {
                background-color: #d4edda;
                color: #155724;
                border-left: 4px solid #28a745;
            }

            .alert-danger {
                background-color: #f8d7da;
                color: #721c24;
                border-left: 4px solid #dc3545;
            }

            .nav-tabs {
                border: none;
                background-color: var(--primary-color);
                padding: 1rem;
                border-radius: 8px;
                margin-bottom: 2rem;
            }

            .nav-tabs .nav-link {
                color: var(--white);
                border: none;
                padding: 1rem 2rem;
                font-weight: 600;
                border-radius: 6px;
                margin-right: 0.5rem;
            }

            .nav-tabs .nav-link.active {
                background-color: var(--white);
                color: var(--primary-color);
            }

            .nav-tabs .nav-link:hover {
                background-color: var(--secondary-color);
            }

            .table {
                background-color: var(--white);
                border-radius: 8px;
                overflow: hidden;
            }

            .table thead th {
                background-color: var(--primary-color);
                color: var(--white);
                border: none;
            }

            .footer {
                background-color: var(--primary-color);
                color: var(--white);
                padding: 1.5rem;
                text-align: center;
                border-radius: 8px;
                margin-top: 2rem;
            }

            .sticky-summary {
                position: sticky;
                top: 20px;
                max-height: calc(100vh - 40px);
                overflow-y: auto;
            }

            .sticky-summary::-webkit-scrollbar {
                width: 8px;
            }

            .sticky-summary::-webkit-scrollbar-track {
                background: #f1f1f1;
                border-radius: 4px;
            }

            .sticky-summary::-webkit-scrollbar-thumb {
                background: var(--primary-color);
                border-radius: 4px;
            }

            .sticky-summary::-webkit-scrollbar-thumb:hover {
                background: var(--secondary-color);
            }
        </style>
    </head>
    <body>
        {%app_entry%}
        <footer>
            {%config%}
            {%scripts%}
            {%renderer%}
        </footer>
    </body>
</html>
'''

# Layout principal
app.layout = dbc.Container([
    # Store para mantener el estado
    dcc.Store(id='session-state', data={
        'bruto': 0,
        'neto': 0,
        'monto': 0,
        'cuotas': 0,
        'tasa': 0,
        'cuota': 0
    }),

    # Navbar
    dbc.Navbar(
        dbc.Container([
            html.Div([
                html.H3("Simulador de Adelanto de Haberes", className="navbar-title", style={"color": "white"})
            ])
        ]),
        color="primary",
        dark=True,
        className="mb-4"
    ),

    # Contenido principal
    dbc.Row([
        # Columna principal (izquierda)
        dbc.Col([
            # Sección 1: Carga de Recibo
            dbc.Card([
                dbc.CardHeader([
                    html.H4("1. Carga de Recibo de Sueldo", className="mb-0 fw-bold"),
                    html.Small("Primero, cargue su recibo de sueldo para continuar", className="text-white")
                ]),
                dbc.CardBody([
                    dcc.Upload(
                        id='upload-pdf',
                        children=html.Div([
                            html.I(className="fas fa-file-pdf me-2"),
                            'Arrastre y suelte o ',
                            html.A('seleccione un archivo PDF', className="text-white fw-bold")
                        ]),
                        style={
                            'width': '100%',
                            'height': '80px',
                            'lineHeight': '80px',
                            'borderWidth': '2px',
                            'borderStyle': 'dashed',
                            'borderRadius': '8px',
                            'textAlign': 'center',
                            'margin': '10px',
                            'backgroundColor': '#2c6aa0',
                            'color': 'white',
                            'cursor': 'pointer',
                            'transition': 'all 0.3s ease'
                        },
                        multiple=False
                    ),
                    html.Div(id='output-pdf-upload')
                ])
            ], className="mb-4"),

            # Sección 2: Simulación
            dbc.Card([
                dbc.CardHeader([
                    html.H4("2. Adelanto de Haberes en Cuotas", className="mb-0 fw-bold"),
                    html.Small("Ingrese los parámetros del adelanto", className="text-white")
                ]),
                dbc.CardBody([
                    dbc.Row([
                        dbc.Col([
                            html.Div([
                                dbc.Input(
                                    id="monto-input",
                                    type="text",
                                    placeholder="Monto solicitado ($)",
                                    className="mb-3"
                                ),
                                # Agregar script para manejar el formateo visual
                                html.Script('''
                                    document.getElementById('monto-input').addEventListener('input', function(e) {
                                        let value = e.target.value.replace(/[^0-9.]/g, '');
                                        if (value) {
                                            let parts = value.split('.');
                                            parts[0] = parts[0].replace(/\B(?=(\d{3})+(?!\d))/g, ',');
                                            e.target.value = '$' + parts.join('.');
                                        }
                                    });
                                ''')
                            ]),
                            dbc.Select(
                                id="cuotas-input",
                                options=[{"label": f"{i} cuotas", "value": i} for i in range(1, 25)],
                                placeholder="Seleccione cantidad de cuotas",
                                className="mb-3"
                            ),
                            html.Div([
                                html.P(f"Tasa anual: {TASA_ANUAL}%", className="mb-2 fw-bold"),
                                html.Small("Tasa fija establecida por el sistema", className="text-muted")
                            ], className="mb-3"),
                            dcc.DatePickerSingle(
                                id="fecha-input",
                                className="mb-3",
                                display_format="DD/MM/YYYY",
                                placeholder="Fecha"
                            ),
                            dbc.Button("Simular", id="simular-button", color="primary", className="mt-3 fw-bold w-100")
                        ], width=6),
                        dbc.Col([
                            html.Div(id="validaciones-simulacion")
                        ], width=6)
                    ]),
                    html.Div(id="simulacion-output")
                ])
            ], className="mb-4"),

            # Sección 3: Generación de Nota
            dbc.Card([
                dbc.CardHeader([
                    html.H4("3. Generación de Nota", className="mb-0 fw-bold"),
                    html.Small("Complete los datos y genere la nota de solicitud", className="text-white")
                ]),
                dbc.CardBody([
                    dbc.Row([
                        dbc.Col([
                            dbc.Input(
                                id="nombre-input",
                                type="text",
                                placeholder="Nombre completo",
                                className="mb-3"
                            ),
                            dbc.Input(
                                id="area-input",
                                type="text",
                                placeholder="Área",
                                className="mb-3"
                            ),
                            dbc.Input(
                                id="sector-input",
                                type="text",
                                placeholder="Sector",
                                className="mb-3"
                            ),
                            dbc.Select(
                                id="motivo-select",
                                options=[{"label": m, "value": m} for m in MOTIVOS],
                                placeholder="Motivo",
                                className="mb-3"
                            ),
                            dbc.Textarea(
                                id="motivo-detallado-input",
                                placeholder="Motivo de la solicitud",
                                className="mb-3"
                            ),
                            dbc.Input(
                                id="puesto-input",
                                type="text",
                                placeholder="Puesto",
                                className="mb-3"
                            ),
                            dbc.Button("Generar Nota", id="generar-nota-button", color="primary", className="mt-3 fw-bold w-100"),
                            html.Div(id="nota-output"),
                            html.Div(id="nota-download", className="mt-3")
                        ], width=6),
                        dbc.Col([
                            html.Div(id="validaciones-nota")
                        ], width=6)
                    ])
                ])
            ])
        ], width=9),

        # Columna lateral (derecha)
        dbc.Col([
            dbc.Card([
                dbc.CardHeader(html.H5("Resumen", className="mb-0 fw-bold")),
                dbc.CardBody([
                    html.Div(id="resumen-sueldo"),
                    html.Hr(),
                    html.Div(id="resumen-prestamo"),
                    html.Hr(),
                    html.Div(id="resumen-nota")
                ])
            ], className="sticky-summary")
        ], width=3)
    ]),

    # Footer
    html.Footer([
        html.P("Sistema de Adelantos Haberes © 2024", className="mb-0")
    ], className="footer")
], fluid=True)

# Callback para actualizar el resumen
@app.callback(
    [Output('resumen-sueldo', 'children'),
     Output('resumen-prestamo', 'children'),
     Output('resumen-nota', 'children')],
    [Input('session-state', 'data'),
     Input('nombre-input', 'value'),
     Input('motivo-select', 'value'),
     Input('motivo-detallado-input', 'value')]
)
def update_resumen(state, nombre, motivo, motivo_detallado):
    resumen_sueldo = []
    resumen_prestamo = []
    resumen_nota = []

    # Resumen de sueldo
    if state.get('bruto', 0) > 0:
        resumen_sueldo = [
            html.H6("Datos del Recibo", className="fw-bold"),
            html.P([
                html.Span("Solicitante: ", className="fw-bold"),
                nombre or state.get('nombre', 'No especificado')
            ]),
            html.P([
                html.Span("Bruto: ", className="fw-bold"),
                f"${state.get('bruto', 0):,.2f}"
            ]),
            html.P([
                html.Span("Neto: ", className="fw-bold"),
                f"${state.get('neto', 0):,.2f}"
            ])
        ]

    # Resumen del préstamo
    if state.get('monto', 0) > 0:
        resumen_prestamo = [
            html.H6("Datos del Préstamo", className="fw-bold"),
            html.P([
                html.Span("Monto: ", className="fw-bold"),
                f"${state.get('monto', 0):,.2f}"
            ]),
            html.P([
                html.Span("Cuotas: ", className="fw-bold"),
                str(state.get('cuotas', 0))
            ]),
            html.P([
                html.Span("Cuota mensual: ", className="fw-bold"),
                f"${state.get('cuota', 0):,.2f}"
            ])
        ]

    # Resumen de la nota
    if motivo:
        resumen_nota = [
            html.H6("Datos de la Solicitud", className="fw-bold"),
            html.P([
                html.Span("Motivo: ", className="fw-bold"),
                motivo
            ])
        ]
        if motivo_detallado:
            resumen_nota.append(
                html.P([
                    html.Span("Motivo detallado: ", className="fw-bold"),
                    motivo_detallado
                ])
            )

    return resumen_sueldo, resumen_prestamo, resumen_nota

# Mantener solo el callback principal que maneja todo
@app.callback(
    [Output('session-state', 'data'),
     Output('output-pdf-upload', 'children'),
     Output('validaciones-simulacion', 'children'),
     Output('nombre-input', 'value'),
     Output('monto-input', 'value'),
     Output('simular-button', 'disabled')],
    [Input('upload-pdf', 'contents'),
     Input('monto-input', 'value'),
     Input('cuotas-input', 'value')],
    [State('upload-pdf', 'filename'),
     State('session-state', 'data')]
)
def update_state_and_outputs(contents, monto_str, cuotas, filename, state):
    ctx = dash.callback_context
    if not ctx.triggered:
        return state, None, None, None, None, True
    
    trigger_id = ctx.triggered[0]['prop_id'].split('.')[0]
    simular_disabled = True
    state = dict(state) if state else {}
    
    if trigger_id == 'monto-input':
        if not monto_str:
            return state, None, None, state.get('nombre', None), None, True
        try:
            monto_limpio = monto_str.replace('$', '').replace(',', '')
            monto_float = float(monto_limpio)
            state['monto'] = monto_float
            log_user_action("MONTO INGRESADO", f"Usuario: {state.get('nombre', 'No especificado')} - Monto: ${monto_float:,.2f} - Tope máximo permitido: ${TOPE_MAXIMO_PRESTAMO:,.2f}")
            log_metric('monto_ingresado', {
                'monto': monto_float,
                'usuario': state.get('nombre', 'No especificado')
            })
            return state, None, None, state.get('nombre', None), monto_str, True
        except:
            return state, None, None, state.get('nombre', None), monto_str, True
    elif trigger_id == 'upload-pdf':
        if contents is None:
            return state, None, None, None, None, True
        try:
            content_type, content_string = contents.split(',')
            decoded = base64.b64decode(content_string)
            with open("temp.pdf", "wb") as f:
                f.write(decoded)
            resultado = calcular_bloques_forzado("temp.pdf")
            if resultado is None:
                log_user_action("ERROR PDF", f"Archivo: {filename} - Error: No se pudieron extraer los datos")
                log_metric('pdf_error', {
                    'filename': filename,
                    'error': 'No se pudieron extraer los datos'
                })
                return state, dbc.Alert("No se pudieron extraer los datos del PDF. Por favor, intente nuevamente.", color="danger"), None, None, None, True
            bruto, deducciones, neto, detectados = resultado
            _, _, nombre_detectado = extraer_sueldos("temp.pdf")
            if bruto is not None and neto is not None:
                state['bruto'] = bruto
                state['neto'] = neto
                state['nombre'] = nombre_detectado
                logging.info(f"PDF subido por: {nombre_detectado} | Bruto: {bruto} | Neto: {neto}")
                log_user_action("PDF PROCESADO", f"Usuario: {nombre_detectado} - Bruto: ${bruto:,.2f} - Neto: ${neto:,.2f}")
                log_metric('pdf_procesado', {
                    'filename': filename,
                    'nombre': nombre_detectado,
                    'bruto': bruto,
                    'neto': neto,
                    'deducciones': deducciones,
                    'conceptos_detectados': detectados
                })
                return state, dbc.Alert([
                    html.H5("Datos extraídos correctamente"),
                    html.P(f"Sueldo bruto: ${bruto:,.2f}"),
                    html.P(f"Sueldo neto: ${neto:,.2f}")
                ], color="success"), None, nombre_detectado, None, True
            else:
                log_user_action("ERROR PDF", f"Archivo: {filename} - Error: Datos incompletos")
                log_metric('pdf_error', {
                    'filename': filename,
                    'error': 'Datos incompletos'
                })
                return state, dbc.Alert("No se pudieron extraer los datos del PDF. Por favor, intente nuevamente.", color="danger"), None, None, None, True
        except Exception as e:
            log_user_action("ERROR PDF", f"Archivo: {filename} - Error: {str(e)}")
            log_metric('pdf_error', {
                'filename': filename,
                'error': str(e)
            })
            print(f"Error al procesar PDF: {str(e)}")
            return state, dbc.Alert(f"Error al procesar el archivo: {str(e)}", color="danger"), None, None, None, True
    elif trigger_id in ['cuotas-input']:
        if monto_str is None or cuotas is None:
            return state, None, None, state.get('nombre', None), state.get('monto', None), True
        validaciones = []
        monto = state.get('monto', 0)
        cuotas = int(cuotas)
        simular_disabled = False  # Inicialmente habilitado
        
        if monto > 3 * state.get('bruto', 0):
            validaciones.append(
                dbc.Alert("El monto excede 3 veces el sueldo bruto.", color="danger")
            )
            simular_disabled = True
            log_user_action("VALIDACIÓN ERROR", f"Usuario: {state.get('nombre', 'No especificado')} - Error: Monto excede 3 veces sueldo bruto - Monto: ${monto:,.2f} - Sueldo: ${state.get('bruto', 0):,.2f} - Tasa anual: {TASA_ANUAL}%")
            log_metric('validacion_error', {
                'tipo': 'tope_sueldo',
                'monto': monto,
                'sueldo_bruto': state.get('bruto', 0),
                'usuario': state.get('nombre', 'No especificado')
            })
        else:
            validaciones.append(
                dbc.Alert("El monto está dentro de los límites permitidos.", color="success")
            )
        try:
            cuota = calcular_cuota(monto, cuotas, TASA_ANUAL)
            if cuota > 0.3 * state.get('neto', 0):
                validaciones.append(
                    dbc.Alert("La cuota mensual excede el 30% del sueldo neto.", color="danger")
                )
                simular_disabled = True
                log_user_action("VALIDACIÓN ERROR", f"Usuario: {state.get('nombre', 'No especificado')} - Error: Cuota excede 30% sueldo neto - Cuota: ${cuota:,.2f} - Sueldo: ${state.get('neto', 0):,.2f} - Tasa anual: {TASA_ANUAL}%")
                log_metric('validacion_error', {
                    'tipo': 'tope_cuota_30',
                    'cuota': cuota,
                    'sueldo_neto': state.get('neto', 0),
                    'usuario': state.get('nombre', 'No especificado')
                })
            elif cuota > TOPE_MAXIMO_PRESTAMO:
                validaciones.append(
                    dbc.Alert(f"La cuota mensual excede el tope máximo permitido de ${TOPE_MAXIMO_PRESTAMO:,.2f}.", color="danger")
                )
                simular_disabled = True
                log_user_action("VALIDACIÓN ERROR", f"Usuario: {state.get('nombre', 'No especificado')} - Error: Cuota excede tope máximo - Cuota: ${cuota:,.2f} - Tope: ${TOPE_MAXIMO_PRESTAMO:,.2f} - Tasa anual: {TASA_ANUAL}%")
                log_metric('validacion_error', {
                    'tipo': 'tope_cuota_max',
                    'cuota': cuota,
                    'tope': TOPE_MAXIMO_PRESTAMO,
                    'usuario': state.get('nombre', 'No especificado')
                })
            else:
                validaciones.append(
                    dbc.Alert(f"Cuota mensual estimada: ${cuota:,.2f}", color="success")
                )
                logging.info(f"Simulación válida: monto={monto}, cuotas={cuotas}, cuota mensual={cuota}")
                log_user_action("SIMULACIÓN VÁLIDA", f"Usuario: {state.get('nombre', 'No especificado')} - Monto: ${monto:,.2f} - Cuotas: {cuotas} - Cuota: ${cuota:,.2f} - Tasa anual: {TASA_ANUAL}% - Tope máximo: ${TOPE_MAXIMO_PRESTAMO:,.2f}")
                log_metric('simulacion_valida', {
                    'monto': monto,
                    'cuotas': cuotas,
                    'cuota': cuota,
                    'tasa': TASA_ANUAL,
                    'usuario': state.get('nombre', 'No especificado')
                })
        except Exception as e:
            print(f"Error al calcular la cuota: {str(e)}")
            validaciones.append(
                dbc.Alert("Error al calcular la cuota mensual.", color="danger")
            )
            cuota = 0
            simular_disabled = True
            log_user_action("ERROR CÁLCULO", f"Usuario: {state.get('nombre', 'No especificado')} - Error: {str(e)} - Tasa anual: {TASA_ANUAL}%")
            log_metric('error_calculo', {
                'error': str(e),
                'monto': monto,
                'cuotas': cuotas,
                'usuario': state.get('nombre', 'No especificado')
            })
        state['cuotas'] = cuotas
        state['tasa'] = TASA_ANUAL
        state['cuota'] = cuota
        return state, None, validaciones, state.get('nombre', None), state.get('monto', None), simular_disabled
    return state, None, None, state.get('nombre', None), state.get('monto', None), True

@app.callback(
    Output('simulacion-output', 'children'),
    Input('simular-button', 'n_clicks'),
    [State('monto-input', 'value'),
     State('cuotas-input', 'value'),
     State('fecha-input', 'date'),
     State('session-state', 'data')]
)
def update_simulacion(n_clicks, monto_str, cuotas, fecha, state):
    if n_clicks is None:
        return None
    print(f"Valores recibidos:")
    print(f"Monto: {monto_str}")
    print(f"Cuotas: {cuotas}")
    print(f"Fecha: {fecha}")
    monto = state.get('monto', 0)
    if monto <= 0:
        return dbc.Alert("Por favor ingrese un monto válido mayor a cero.", color="danger")
    try:
        cuotas = int(cuotas) if cuotas is not None else None
    except (ValueError, TypeError):
        print("Error al convertir valores")
        return dbc.Alert("Error en los valores ingresados. Por favor, verifique los datos.", color="danger")
    campos_faltantes = []
    if cuotas is None or cuotas <= 0:
        campos_faltantes.append("Cuotas")
    if not fecha:
        campos_faltantes.append("Fecha")
    if campos_faltantes:
        return dbc.Alert(
            f"Por favor complete los siguientes campos: {', '.join(campos_faltantes)}",
            color="danger"
        )
    try:
        df_amort = generar_cuadro_amortizacion(monto, cuotas, TASA_ANUAL)
        logging.info(f"Simulación realizada: monto={monto}, cuotas={cuotas}, fecha={fecha}")
        log_user_action("SIMULACIÓN REALIZADA", f"Usuario: {state.get('nombre', 'No especificado')} - Monto: ${monto:,.2f} - Cuotas: {cuotas} - Fecha: {fecha} - Cuota mensual: ${calcular_cuota(monto, cuotas, TASA_ANUAL):,.2f} - Tasa anual: {TASA_ANUAL}% - Tope máximo: ${TOPE_MAXIMO_PRESTAMO:,.2f}")
        return [
            html.H4("Resumen de la simulación"),
            dbc.Row([
                dbc.Col([
                    html.P(f"Monto solicitado: ${monto:,.2f}"),
                    html.P(f"Cantidad de cuotas: {cuotas}"),
                    html.P(f"Cuota mensual estimada: ${calcular_cuota(monto, cuotas, TASA_ANUAL):,.2f}")
                ], width=6),
                dbc.Col([
                    html.P(f"Tasa anual: {TASA_ANUAL:.2f}%"),
                    html.P(f"Tasa mensual: {TASA_ANUAL/12:.2f}%"),
                    html.P(f"Fecha de inicio: {fecha}")
                ], width=6)
            ]),
            html.H4("Cuadro de Amortización"),
            dbc.Table.from_dataframe(
                df_amort,
                striped=True,
                bordered=True,
                hover=True
            )
        ]
    except Exception as e:
        print(f"Error en la simulación: {str(e)}")
        log_user_action("ERROR SIMULACIÓN", f"Usuario: {state.get('nombre', 'No especificado')} - Error: {str(e)} - Tasa anual: {TASA_ANUAL}% - Tope máximo: ${TOPE_MAXIMO_PRESTAMO:,.2f}")
        return dbc.Alert(f"Error al generar la simulación: {str(e)}", color="danger")

# Diccionario temporal para guardar archivos generados por sesión
GENERATED_FILES = {}

@app.server.route('/download/<file_id>')
def download_file(file_id):
    file_path = GENERATED_FILES.get(file_id)
    if file_path and os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return "Archivo no encontrado", 404

@app.callback(
    Output('nota-output', 'children'),
    Output('nota-download', 'children'),
    Input('generar-nota-button', 'n_clicks'),
    [State('nombre-input', 'value'),
     State('area-input', 'value'),
     State('sector-input', 'value'),
     State('motivo-select', 'value'),
     State('motivo-detallado-input', 'value'),
     State('puesto-input', 'value'),
     State('session-state', 'data')]
)
def generar_nota_callback(n_clicks, nombre, area, sector, motivo, motivo_detallado, puesto, state):
    if n_clicks is None:
        return None, None
    if not all([nombre, area, sector, motivo, motivo_detallado, puesto]):
        log_user_action("ERROR NOTA", f"Usuario: {nombre or 'No especificado'} - Error: Datos incompletos")
        log_metric('nota_error', {
            'error': 'Datos incompletos',
            'datos_proporcionados': {
                'nombre': bool(nombre),
                'area': bool(area),
                'sector': bool(sector),
                'motivo': bool(motivo),
                'motivo_detallado': bool(motivo_detallado),
                'puesto': bool(puesto)
            }
        })
        return dbc.Alert("Por favor complete todos los datos del usuario.", color="danger"), None
    
    logging.info(f"Nota generada para: {nombre} | Motivo: {motivo} | Detalle: {motivo_detallado} | Área: {area} | Sector: {sector} | Puesto: {puesto}")
    log_user_action("NOTA GENERADA", f"Usuario: {nombre} - Área: {area} - Sector: {sector} - Motivo: {motivo} - Monto: ${state.get('monto', 0):,.2f}")
    log_metric('nota_generada', {
        'nombre': nombre,
        'area': area,
        'sector': sector,
        'motivo': motivo,
        'motivo_detallado': motivo_detallado,
        'puesto': puesto,
        'monto': state.get('monto', 0),
        'cuotas': state.get('cuotas', 0),
        'tasa': state.get('tasa', 0),
        'cuota': state.get('cuota', 0),
        'sueldo_neto': state.get('neto', 0)
    })
    
    docx_bytes = generar_nota(
        state.get('monto', 0),
        state.get('cuotas', 0),
        state.get('tasa', 0),
        state.get('cuota', 0),
        datetime.now(),
        nombre, area, sector, motivo, motivo_detallado, puesto,
        state.get('neto', 0)
    )
    if docx_bytes is not None:
        temp_dir = tempfile.gettempdir()
        file_id = str(uuid.uuid4())
        file_path = os.path.join(temp_dir, f"nota_{file_id}.docx")
        with open(file_path, "wb") as f:
            f.write(docx_bytes.getvalue())
        GENERATED_FILES[file_id] = file_path
        href = f"/download/{file_id}"
        return (
            dbc.Alert("✅ Nota generada correctamente.", color="success"),
            html.A(
                "Descargar Nota de Solicitud",
                href=href,
                className="btn btn-primary w-100"
            )
        )
    else:
        log_user_action("ERROR NOTA", f"Usuario: {nombre} - Error: No se pudo generar el archivo")
        log_metric('nota_error', {
            'error': 'Error al generar el archivo',
            'datos': {
                'nombre': nombre,
                'area': area,
                'sector': sector,
                'motivo': motivo
            }
        })
        return dbc.Alert("❌ No se pudo generar la nota. Por favor, intente nuevamente.", color="danger"), None

# Mantener las funciones auxiliares existentes
def extraer_sueldos(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()

        lines = text.splitlines()
        
        # Extraer nombre (buscando después de "Apellido y Nombre:")
        nombre = None
        for i, line in enumerate(lines):
            if "Apellido y Nombre:" in line:
                # Buscar la siguiente línea que contenga una coma (formato "Apellido, Nombre")
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if "," in next_line and not any(field in next_line for field in ["Categoria:", "Cargo:", "Egreso:", "Codigo", "Concepto"]):
                        nombre = next_line
                        # Convertir de "Apellido, Nombre" a "Nombre Apellido"
                        if "," in nombre:
                            apellido, nombre_persona = nombre.split(",", 1)
                            nombre = f"{nombre_persona.strip()} {apellido.strip()}"
                        break
                    j += 1
                break
        
        # Extraer montos
        monto_regex = re.compile(r'^\s*(\d{1,3}(?:\.\d{3})*,\d{2})\s*$')
        valores = [match.group(1) for line in lines if (match := monto_regex.match(line))]

        if len(valores) < 2:
            print("No se encontraron suficientes montos claros para bruto/neto")
            return None, None, None

        valores_f = [float(v.replace('.', '').replace(',', '.')) for v in valores]
        sueldo_neto = valores_f[-1]
        candidatos = [v for v in valores_f[-6:] if v > 1_000_000]
        if not candidatos:
            print("No se detectó un valor alto para el sueldo bruto.")
            return None, None, None
        sueldo_bruto = max(candidatos)

        return sueldo_bruto, sueldo_neto, nombre

    except Exception as e:
        print(f"Error al procesar PDF: {e}")
        return None, None, None

def calcular_bloques_forzado(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()

        lines = text.splitlines()
        bruto = 0.0
        deducciones = 0.0
        detectados = []

        def es_monto(s):
            return re.match(r'^-?\d{1,3}(?:\.\d{3})*,\d{2}$', s)
        def es_cantidad(s):
            return re.match(r'^\d{1,3}(?:\.\d{3})*,\d{2}$', s)

        # Buscar la sección de conceptos
        inicio_conceptos = False
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            # Detectar inicio de la sección de conceptos
            if line == "Codigo":
                inicio_conceptos = True
                i += 1
                continue
            if inicio_conceptos:
                for codigo in CODIGOS_BRUTO.keys():
                    # Solo considerar líneas que empiezan por el código, espacio y una letra (no número ni coma)
                    if re.match(rf'^{codigo} [A-Za-z]', line):
                        # Caso 1: cantidad y luego monto
                        if i + 2 < len(lines) and es_cantidad(lines[i+1].strip()) and es_monto(lines[i+2].strip()):
                            valor_str = lines[i+2].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            bruto += valor
                            detectados.append((codigo, valor, "REM", line))
                        # Caso 2: monto directo
                        elif i + 1 < len(lines) and es_monto(lines[i+1].strip()):
                            valor_str = lines[i+1].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            bruto += valor
                            detectados.append((codigo, valor, "REM", line))
                        break
                for codigo in CODIGOS_DEDUCCIONES.keys():
                    if re.match(rf'^{codigo} [A-Za-z]', line):
                        if i + 2 < len(lines) and es_cantidad(lines[i+1].strip()) and es_monto(lines[i+2].strip()):
                            valor_str = lines[i+2].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            deducciones += valor
                            detectados.append((codigo, valor, "DED", line))
                        elif i + 1 < len(lines) and es_monto(lines[i+1].strip()):
                            valor_str = lines[i+1].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            deducciones += valor
                            detectados.append((codigo, valor, "DED", line))
                        break
            i += 1

        neto = bruto - deducciones
        return round(bruto, 2), round(deducciones, 2), round(neto, 2), detectados

    except Exception as e:
        print(f"Error al procesar PDF: {e}")
        return None

def calcular_cuota(monto, cuotas, tasa_anual):
    if monto is None or cuotas is None or tasa_anual is None:
        return None
        
    # Asegurarse de que los valores sean del tipo correcto
    monto = float(monto)
    cuotas = int(cuotas)
    tasa_anual = float(tasa_anual)
        
    tasa_mensual = (tasa_anual / 100) / 12
    if tasa_mensual == 0:
        return monto / cuotas
    cuota = monto * (tasa_mensual * (1 + tasa_mensual)**cuotas) / ((1 + tasa_mensual)**cuotas - 1)
    return round(cuota, 2)

def generar_cuadro_amortizacion(monto, cuotas, tasa_anual):
    if monto is None or cuotas is None or tasa_anual is None:
        return pd.DataFrame()
        
    tasa_mensual = (tasa_anual / 100) / 12
    cuota_total = calcular_cuota(monto, cuotas, tasa_anual)
    saldo = monto
    cuadro = []
    for i in range(1, cuotas + 1):
        interes = saldo * tasa_mensual
        amortizacion = cuota_total - interes
        saldo -= amortizacion
        cuadro.append({
            "Cuota N°": i,
            "Cuota total ($)": round(cuota_total, 2),
            "Interés ($)": round(interes, 2),
            "Amortización ($)": round(amortizacion, 2),
            "Saldo restante ($)": round(saldo if saldo > 0 else 0, 2)
        })
    return pd.DataFrame(cuadro)

def monto_a_letras_bancario(monto):
    entero = int(monto)
    decimales = int(round((monto - entero) * 100))
    texto = num2words(entero, lang='es').replace("uno", "un").capitalize()
    if decimales == 0:
        return f"{texto} pesos"
    else:
        texto_centavos = num2words(decimales, lang='es').replace("uno", "un")
        return f"{texto} pesos con {texto_centavos} centavos"

def generar_nota(monto, cuotas, tasa_final, cuota, fecha, nombre, area, sector, motivo, motivo_detallado, puesto, neto):
    def formatear_fecha_larga(fecha):
        meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                 'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
        return f"{fecha.day} de {meses[fecha.month - 1]} del {fecha.year}"

    def tercer_viernes(fecha_base):
        year = fecha_base.year
        month = fecha_base.month
        count = 0
        for day in range(1, 32):
            try:
                fecha = datetime(year, month, day)
                if fecha.weekday() == 4:
                    count += 1
                    if count == 3:
                        return fecha
            except ValueError:
                break
        return fecha_base

    def ultimo_dia_habil_del_mes(fecha_base):
        anio, mes = fecha_base.year, fecha_base.month
        ultimo_dia = calendar.monthrange(anio, mes)[1]
        venc = datetime(anio, mes, ultimo_dia)
        while venc.weekday() >= 5:
            venc -= timedelta(days=1)
        return venc

    try:
        fecha_directorio = tercer_viernes(fecha)
        vencimiento = ultimo_dia_habil_del_mes(fecha)
        texto_letras = monto_a_letras_bancario(monto)
        neto_menos_cuota = neto - cuota
        neto_menos_cuota_letras = monto_a_letras_bancario(neto_menos_cuota)
        datos = {
            "<nombre>": nombre,
            "<area>": area,
            "<sector>": sector,
            "<fecha>": formatear_fecha_larga(fecha),
            "<fecha_directorio>": formatear_fecha_larga(fecha_directorio),
            "<monto>": f"${monto:,.2f}",
            "<cuotas>": str(cuotas),
            "<cuotas_en_letras>": num2words(cuotas, lang='es').replace("uno", "un").capitalize(),
            "<motivo>": motivo,
            "<detalle_motivo>": motivo_detallado,
            "<monto_en_letras>": texto_letras,
            "<tasa>": f"{tasa_final:.2f}%",
            "<vencimiento>": formatear_fecha_larga(vencimiento),
            "<puesto>": puesto,
            "<neto_menos_cuota>": f"${neto_menos_cuota:,.2f}",
            "<neto_menos_cuota_letras>": neto_menos_cuota_letras
        }

        print("Diccionario de datos:")
        for k, v in datos.items():
            print(f"{k}: {v}")

        plantilla = None
        for archivo in os.listdir(os.getcwd()):
            if archivo.endswith(".docx") and "nota" in archivo.lower():
                doc_test = Document(archivo)
                texts = [p.text for p in doc_test.paragraphs]
                texts += [c.text for t in doc_test.tables for r in t.rows for c in r.cells]
                if any("<" in t and ">" in t for t in texts):
                    plantilla = archivo
                    break

        if not plantilla:
            print("No se encontró una plantilla con '<>' en la carpeta.")
            return None

        print(f"Plantilla seleccionada: {plantilla}")
        print(f"Ruta absoluta: {os.path.abspath(plantilla)}")
        doc = Document(plantilla)

        for p in doc.paragraphs:
            for k, v in datos.items():
                if k in p.text:
                    for r in p.runs:
                        r.text = r.text.replace(k, v)
                        # Aplicar formato Garamond 13pt a todos los reemplazos
                        r.font.name = 'Garamond'
                        r.font.size = Pt(13)
                if k in p.text:
                    p.text = p.text.replace(k, v)

        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for k, v in datos.items():
                            if k in p.text:
                                for run in p.runs:
                                    run.text = run.text.replace(k, v)
                                    # Aplicar formato Garamond 13pt a todos los reemplazos
                                    run.font.name = 'Garamond'
                                    run.font.size = Pt(13)
                            if k in p.text:
                                p.text = p.text.replace(k, v)

        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for k, v in datos.items():
                        if k in c.text:
                            c.text = c.text.replace(k, v)
                            # Aplicar formato Garamond 13pt en celdas de tabla
                            for p in c.paragraphs:
                                for run in p.runs:
                                    run.font.name = 'Garamond'
                                    run.font.size = Pt(13)

        try:
            df_amort = generar_cuadro_amortizacion(monto, cuotas, tasa_final)
            for i, p in enumerate(doc.paragraphs):
                if "<cuadro_amortizacion>" in p.text:
                    p.text = p.text.replace("<cuadro_amortizacion>", "")
                    table = doc.add_table(rows=1, cols=len(df_amort.columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, col in enumerate(df_amort.columns):
                        hdr_cells[j].text = str(col)
                    for _, row in df_amort.iterrows():
                        row_cells = table.add_row().cells
                        for j, val in enumerate(row):
                            row_cells[j].text = str(val)
                    p._p.addnext(table._tbl)
                    break
        except Exception as e:
            print(f"No se pudo insertar la tabla: {e}")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes

    except Exception as e:
        print(f"Error al generar nota: {e}")
        return None

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8050))
    app.run_server(host="0.0.0.0", port=port)
